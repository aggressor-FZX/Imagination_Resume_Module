#!/usr/bin/env python3
"""Build a resume PDF from an Imaginator export bundle and a LaTeX template.

The script accepts either a raw export-results text file or a JSON file that
contains the structured objects emitted by the Imaginator pipeline. It extracts
the most relevant data, maps it into the section structure used by
main_template.tex, writes a generated .tex file, and compiles it to PDF.
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import tempfile
import base64
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple, Tuple


DEFAULT_TEMPLATE = Path(__file__).parent / "main_template.tex"
DEFAULT_OUTPUT_DIR = Path("/app/data/generated_resumes")


@dataclass
class ResumeSection:
    title: str
    body_lines: List[str] = field(default_factory=list)
    start_date: str = ""
    end_date: str = ""
    date_string: str = ""  # pre-formatted date range if provided directly


@dataclass
class ResumePayload:
    name: str
    location: str = ""
    phone: str = ""
    email: str = ""
    linkedin: str = ""
    github: str = ""
    headline: str = ""
    summary: str = ""
    skills: List[Tuple[str, List[str]]] = field(default_factory=list)
    experience: List[ResumeSection] = field(default_factory=list)
    projects: List[ResumeSection] = field(default_factory=list)
    education: List[str] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)
    publication: str = ""


LATEX_SPECIAL_CHARS = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def latex_escape(text: str) -> str:
    escaped = text
    for original, replacement in LATEX_SPECIAL_CHARS.items():
        escaped = escaped.replace(original, replacement)
    return escaped


def normalize_dash(text: str) -> str:
    return text.replace("–", "--").replace("—", "--")


def split_list_block(text: str) -> List[str]:
    items: List[str] = []
    for raw in re.split(r"\s*(?:\||\n|•|\*)\s*", text):
        item = raw.strip().strip("-").strip()
        if item:
            items.append(item)
    return items


def extract_json_block(text: str) -> Optional[Dict[str, Any]]:
    decoder = json.JSONDecoder()
    marker = text.find('"exportedAt"')
    if marker == -1:
        marker = text.find("{\n  \"exportedAt\"")
    if marker == -1:
        marker = text.find("{\n\"exportedAt\"")
    if marker == -1:
        return None
    start = text.rfind("{", 0, marker)
    if start == -1:
        start = marker

    candidate = text[start:]
    for offset in range(len(candidate)):
        try:
            obj, end = decoder.raw_decode(candidate[offset:])
            if isinstance(obj, dict) and "data" in obj:
                return obj
        except json.JSONDecodeError:
            continue
    return None


def get_nested(data: Dict[str, Any], keys: Sequence[str], default: Any = None) -> Any:
    current: Any = data
    for key in keys:
        if isinstance(current, dict) and key in current:
            current = current[key]
        else:
            return default
    return current


def coerce_list(value: Any) -> List[Any]:
    if isinstance(value, list):
        return value
    if value is None:
        return []
    return [value]


def first_nonempty(mapping: Dict[str, Any], keys: Sequence[str], default: str = "") -> str:
    for key in keys:
        value = mapping.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return default


def format_entry_title(entry: Dict[str, Any], fallback: str = "") -> str:
    return first_nonempty(
        entry,
        ["title_line", "title", "role", "name", "project", "label"],
        fallback,
    )


def format_entry_body(entry: Dict[str, Any]) -> List[str]:
    bullets = []
    for key in ["bullets", "highlights", "items", "achievements", "responsibilities"]:
        value = entry.get(key)
        if isinstance(value, list):
            for item in value:
                if isinstance(item, str) and item.strip():
                    bullets.append(item.strip())
        elif isinstance(value, str) and value.strip():
            bullets.extend([line.strip() for line in re.split(r"[\n;]", value) if line.strip()])

    if not bullets:
        for key in ["snippet", "summary", "description", "body", "text"]:
            value = entry.get(key)
            if isinstance(value, str) and value.strip():
                bullets.extend([line.strip() for line in re.split(r"[\n;]", value) if line.strip()])
                break

    return bullets


def format_education_entry(entry: Any) -> str:
    if isinstance(entry, str):
        return entry.strip()
    if not isinstance(entry, dict):
        return str(entry)

    degree = first_nonempty(entry, ["degree", "title", "name"], "")
    field = first_nonempty(entry, ["field", "major"], "")
    institution = first_nonempty(entry, ["institution", "school", "university"], "")
    year = first_nonempty(entry, ["year", "date", "date_completed"], "")
    gpa = first_nonempty(entry, ["gpa"], "")

    parts = [part for part in [degree, field, institution, year] if part]
    line = ", ".join(parts)
    if gpa:
        line = f"{line} | GPA: {gpa}" if line else f"GPA: {gpa}"
    return line or str(entry)


def format_certification_entry(entry: Any) -> str:
    if isinstance(entry, str):
        return entry.strip()
    if not isinstance(entry, dict):
        return str(entry)
    name = first_nonempty(entry, ["certification", "name", "title"], "")
    issuer = first_nonempty(entry, ["issuer", "organization", "company"], "")
    date = first_nonempty(entry, ["date", "year", "issued"], "")
    parts = [part for part in [name, issuer, date] if part]
    return " ".join(parts) or str(entry)


def parse_skill_groups(source: Dict[str, Any]) -> List[Tuple[str, List[str]]]:
    structured = source.get("skills")
    if isinstance(structured, list):
        if structured and all(isinstance(item, str) for item in structured):
            return [("Skills", [item.strip() for item in structured if item.strip()])]
        if structured and all(isinstance(item, dict) for item in structured):
            groups: List[Tuple[str, List[str]]] = []
            for entry in structured:
                label = first_nonempty(entry, ["category", "title", "name", "label"], "Skills")
                values: List[str] = []
                for key in ["items", "skills", "values"]:
                    raw = entry.get(key)
                    if isinstance(raw, list):
                        values.extend([str(item).strip() for item in raw if str(item).strip()])
                if not values:
                    values = [str(entry.get("skill", "")).strip()] if entry.get("skill") else []
                groups.append((label, values))
            return groups

    flat = source.get("user_skills") or source.get("aggregate_skills") or source.get("skills_list")
    if isinstance(flat, list):
        items = [str(item).strip() for item in flat if str(item).strip()]
        if items:
            return [("Skills", items)]

    return []


def format_date_range(entry: Dict[str, Any]) -> str:
    """Extract a date range from an entry dict.
    
    Supports: date_string, start_date/end_date, dates, date, year
    Returns a formatted string like '12/2019 -- 05/2023' or ''.
    """
    # Direct date string
    for key in ["date_string", "dates", "date_range"]:
        val = entry.get(key)
        if isinstance(val, str) and val.strip():
            return val.strip()
    
    start = first_nonempty(entry, ["start_date", "start", "from", "date_start"], "")
    end = first_nonempty(entry, ["end_date", "end", "to", "date_end"], "")
    
    if start and end:
        return f"{start} -- {end}"
    if start:
        return f"{start} -- Present"
    if end:
        return end
    
    # Single date/year fallback
    for key in ["date", "year"]:
        val = entry.get(key)
        if isinstance(val, str) and val.strip():
            return val.strip()
    
    return ""


def parse_sections(source: Dict[str, Any], keys: Sequence[str]) -> List[ResumeSection]:
    """Parse experience/project entries with date support."""
    for key in keys:
        value = source.get(key)
        if isinstance(value, list) and value:
            sections: List[ResumeSection] = []
            for index, entry in enumerate(value):
                if isinstance(entry, dict):
                    title = format_entry_title(entry, f"Entry {index + 1}")
                    body = format_entry_body(entry)
                    if not body and entry.get("text"):
                        body = [str(entry["text"]).strip()]
                    date_str = format_date_range(entry)
                    sections.append(ResumeSection(
                        title=title or f"Entry {index + 1}",
                        body_lines=body,
                        date_string=date_str,
                    ))
                else:
                    sections.append(ResumeSection(
                        title=f"Entry {index + 1}",
                        body_lines=[str(entry).strip()],
                    ))
            return sections
    return []


def extract_text_bundle(path: Path) -> Dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    payload = extract_json_block(text)
    if payload:
        return payload

    raise ValueError(f"Could not locate structured JSON payload in {path}")


def build_fallback_payload(source: Dict[str, Any]) -> ResumePayload:
    data = source.get("data", source)
    rewritten = data.get("rewrittenResume", "") or ""
    domain_insights = data.get("domainInsights", {}) or {}

    name = "Jeff Calderon"
    location = get_nested(domain_insights, ["jobLocation"], "") or data.get("jobLocation", "") or "Marysville, WA"
    phone = "916.712.8396"
    email = "jeff.d.calderon@gmail.com"
    linkedin = "linkedin.com/in/jeffdcalderon"
    github = "github.com/aggressor-FZX"
    headline = data.get("jobTitle", "") or get_nested(domain_insights, ["onet", "title"], "") or "Engineering Data Scientist"

    summary = (
        "Engineering data scientist with a physics foundation and 10+ years of applied experience designing predictive models, "
        "analyzing sensor and temporal data streams, and delivering actionable safety insights to operational and executive stakeholders. "
        "Background spans safety-critical federal research, multi-sensor fleet operations, aviation-adjacent work, and Boeing Defense-sponsored ML research."
    )

    skills = [
        (
            "Statistical and Predictive Modeling",
            [
                "descriptive analysis",
                "predictive modeling",
                "Monte Carlo simulation",
                "anomaly detection",
                "model validation",
                "risk assessment",
            ],
        ),
        (
            "Languages and Tools",
            [
                "Python",
                "SQL",
                "C++",
                "Bash",
                "Pandas",
                "NumPy",
                "PyTorch",
                "TensorFlow",
                "AWS",
                "Docker",
                "Git",
            ],
        ),
        (
            "NLP and AI",
            [
                "LLM integration",
                "LangChain",
                "OpenAI API",
                "RAG",
                "embeddings",
                "prompt engineering",
            ],
        ),
    ]

    experience = [
        ResumeSection(
            title="Data Analyst Intern -- Washington State Data Exchange for Public Safety",
            body_lines=[
                "Design and implement ETL/ELT pipelines (AWS Step Functions, Glue, Lambda) ingesting multi-source safety and operational data; build Power BI dashboards and reports tracking safety-relevant KPIs for law enforcement and EMS leadership.",
                "Perform exploratory analysis and variance reporting on operational datasets; translate findings into actionable recommendations for agency directors and cross-functional stakeholders.",
            ],
        ),
        ResumeSection(
            title="Operations Manager -- National Oceanic and Atmospheric Administration (NOAA)",
            body_lines=[
                "Built descriptive and predictive analytical models synthesizing high-dimensional sensor data from 13 oceanographic research vessels; designed Python-based decision-support tools for operational safety and readiness monitoring, reducing manual reporting overhead by 40%.",
                "Delivered safety-relevant data insights and risk-based recommendations to Regional Director and fleet leadership through regular executive briefings.",
            ],
        ),
        ResumeSection(
            title="Marine Data Analyst | Hydrographer -- NOAA Commissioned Officer Corps (Office of Coast Survey)",
            body_lines=[
                "Processed and fused multi-sensor shipboard data streams in support of the U.S. Office of Coast Survey: sonar, LiDAR, GPS, inertial motion, tidal modeling, and shipboard weather sensor arrays.",
                "Designed and implemented anomaly detection algorithms on sonar and geospatial datasets to identify navigation hazards and seafloor irregularities for official nautical chart updates.",
            ],
        ),
        ResumeSection(
            title="Research Physicist -- Naval Research Laboratory",
            body_lines=[
                "Designed original Monte Carlo algorithms for probabilistic safety risk quantification and radiological threat detection modeling in Python and C++; evaluated algorithm performance against empirical benchmarks and documented findings for reproducibility.",
                "Results secured $500K in federal safety R&D funding; co-authored peer-reviewed publication and operated in a classified national security environment.",
            ],
        ),
    ]

    projects = [
        ResumeSection(
            title="Boeing Defense -- Aerial Safety Vision AI",
            body_lines=[
                "Trained generative models on 263,000+ image datasets via PyTorch DDP on NVIDIA H100/A100 GPU clusters; designed evaluation frameworks benchmarked across real-world flight scenarios.",
            ],
        ),
        ResumeSection(
            title="Sensor Data Analysis -- LHC Particle Physics | University of Maryland",
            body_lines=[
                "Designed and implemented gradient-based solvers, efficiency quantification algorithms, and statistical significance methods to extract signal from multi-terabyte particle physics sensor datasets at CERN.",
            ],
        ),
    ]

    education = [
        "B.S. in Data Analytics, Washington State University, Everett -- 05/2026",
        "A.A.S. in Computer Science, Everett Community College -- 04/2024 -- GPA: 4.0 (President's List)",
        "B.S. in Physics, University of Maryland, College Park -- 12/2015",
    ]

    certifications = [
        "Network Security+, CompTIA (07/2024)",
        "Cybersecurity Analyst, Everett Community College (07/2024)",
        "Data Analytics, Washington State University (12/2024)",
        "Medical Person in Charge, U.S. Coast Guard Certified",
    ]

    publication = "Journal of Instrumentation (2016) -- Liquid Scintillator Tiles for Calorimetry, Large Hadron Collider detector performance analysis"

    return ResumePayload(
        name=name,
        location=location,
        phone=phone,
        email=email,
        linkedin=linkedin,
        github=github,
        headline=headline,
        summary=summary,
        skills=skills,
        experience=experience,
        projects=projects,
        education=education,
        certifications=certifications,
        publication=publication,
    )


def parse_resume_payload(bundle: Dict[str, Any]) -> ResumePayload:
    source = bundle.get("data", bundle)

    # Preferred path: direct structured input from the IMAGINATOR input spec.
    if any(key in source for key in ["experience", "professional_experience", "projects", "education", "certifications", "skills"]):
        name = first_nonempty(source, ["name", "full_name"], "Jeff Calderon")
        location = first_nonempty(source, ["location", "job_location", "preferred_location", "city"], "Marysville, WA")
        phone = first_nonempty(source, ["phone", "phone_number"], "916.712.8396")
        email = first_nonempty(source, ["email", "email_address"], "jeff.d.calderon@gmail.com")
        linkedin = first_nonempty(source, ["linkedin", "linkedin_url"], "linkedin.com/in/jeffdcalderon")
        github = first_nonempty(source, ["github", "github_url"], "github.com/aggressor-FZX")
        headline = first_nonempty(source, ["job_title", "target_job_title", "headline"], "Engineering Data Scientist")
        summary = first_nonempty(source, ["summary", "professional_summary", "profile_summary"], "") or (
            "Engineering data scientist with a physics foundation and 10+ years of applied experience designing predictive models, "
            "analyzing sensor and temporal data streams, and delivering actionable insights to operational and executive stakeholders."
        )

        skills = parse_skill_groups(source)
        experience = parse_sections(source, ["experience", "professional_experience", "experiences"])
        projects = parse_sections(source, ["projects"])
        education_raw = coerce_list(source.get("education"))
        certifications_raw = coerce_list(source.get("certifications"))
        publication = first_nonempty(source, ["publication", "publications", "publication_line"], "")

        if not skills:
            skills = [
                ("Skills", [skill for skill in coerce_list(source.get("user_skills")) if isinstance(skill, str) and skill.strip()]),
            ]
            skills = [item for item in skills if item[1]]

        if not experience:
            fallback_text = source.get("resume_text") or source.get("rewrittenResume") or ""
            if fallback_text:
                experience = [ResumeSection(title="Resume Experience", body_lines=[line.strip() for line in re.split(r"[\n•;]", str(fallback_text)) if line.strip()][:12])]

        if not projects:
            projects = []

        return ResumePayload(
            name=name,
            location=location,
            phone=phone,
            email=email,
            linkedin=linkedin,
            github=github,
            headline=headline,
            summary=summary,
            skills=skills,
            experience=experience or [],
            projects=projects,
            education=[format_education_entry(entry) for entry in education_raw if format_education_entry(entry)],
            certifications=[format_certification_entry(entry) for entry in certifications_raw if format_certification_entry(entry)],
            publication=publication or "",
        )

    # Fallback path: current export bundle sample.
    data = source
    rewritten = data.get("rewrittenResume", "") or ""
    domain_insights = data.get("domainInsights", {}) or {}

    name = "Jeff Calderon"
    location = get_nested(domain_insights, ["jobLocation"], "") or data.get("jobLocation", "") or "Marysville, WA"
    phone = "916.712.8396"
    email = "jeff.d.calderon@gmail.com"
    linkedin = "linkedin.com/in/jeffdcalderon"
    github = "github.com/aggressor-FZX"
    headline = data.get("jobTitle", "") or get_nested(domain_insights, ["onet", "title"], "") or "Engineering Data Scientist"

    summary = (
        "Engineering data scientist with a physics foundation and 10+ years of applied experience designing predictive models, "
        "analyzing sensor and temporal data streams, and delivering actionable safety insights to operational and executive stakeholders. "
        "Background spans safety-critical federal research, multi-sensor fleet operations, aviation-adjacent work, and Boeing Defense-sponsored ML research."
    )

    skills = [
        (
            "Statistical and Predictive Modeling",
            [
                "descriptive analysis",
                "predictive modeling",
                "Monte Carlo simulation",
                "anomaly detection",
                "model validation",
                "risk assessment",
            ],
        ),
        (
            "Languages and Tools",
            [
                "Python",
                "SQL",
                "C++",
                "Bash",
                "Pandas",
                "NumPy",
                "PyTorch",
                "TensorFlow",
                "AWS",
                "Docker",
                "Git",
            ],
        ),
        (
            "NLP and AI",
            [
                "LLM integration",
                "LangChain",
                "OpenAI API",
                "RAG",
                "embeddings",
                "prompt engineering",
            ],
        ),
    ]

    experience = [
        ResumeSection(
            title="Data Analyst Intern -- Washington State Data Exchange for Public Safety",
            body_lines=[
                "Design and implement ETL/ELT pipelines (AWS Step Functions, Glue, Lambda) ingesting multi-source safety and operational data; build Power BI dashboards and reports tracking safety-relevant KPIs for law enforcement and EMS leadership.",
                "Perform exploratory analysis and variance reporting on operational datasets; translate findings into actionable recommendations for agency directors and cross-functional stakeholders.",
            ],
        ),
        ResumeSection(
            title="Operations Manager -- National Oceanic and Atmospheric Administration (NOAA)",
            body_lines=[
                "Built descriptive and predictive analytical models synthesizing high-dimensional sensor data from 13 oceanographic research vessels; designed Python-based decision-support tools for operational safety and readiness monitoring, reducing manual reporting overhead by 40%.",
                "Delivered safety-relevant data insights and risk-based recommendations to Regional Director and fleet leadership through regular executive briefings.",
            ],
        ),
        ResumeSection(
            title="Marine Data Analyst | Hydrographer -- NOAA Commissioned Officer Corps (Office of Coast Survey)",
            body_lines=[
                "Processed and fused multi-sensor shipboard data streams in support of the U.S. Office of Coast Survey: sonar, LiDAR, GPS, inertial motion, tidal modeling, and shipboard weather sensor arrays.",
                "Designed and implemented anomaly detection algorithms on sonar and geospatial datasets to identify navigation hazards and seafloor irregularities for official nautical chart updates.",
            ],
        ),
        ResumeSection(
            title="Research Physicist -- Naval Research Laboratory",
            body_lines=[
                "Designed original Monte Carlo algorithms for probabilistic safety risk quantification and radiological threat detection modeling in Python and C++; evaluated algorithm performance against empirical benchmarks and documented findings for reproducibility.",
                "Results secured $500K in federal safety R&D funding; co-authored peer-reviewed publication and operated in a classified national security environment.",
            ],
        ),
    ]

    projects = [
        ResumeSection(
            title="Boeing Defense -- Aerial Safety Vision AI",
            body_lines=[
                "Trained generative models on 263,000+ image datasets via PyTorch DDP on NVIDIA H100/A100 GPU clusters; designed evaluation frameworks benchmarked across real-world flight scenarios.",
            ],
        ),
        ResumeSection(
            title="Sensor Data Analysis -- LHC Particle Physics | University of Maryland",
            body_lines=[
                "Designed and implemented gradient-based solvers, efficiency quantification algorithms, and statistical significance methods to extract signal from multi-terabyte particle physics sensor datasets at CERN.",
            ],
        ),
    ]

    education = [
        "B.S. in Data Analytics, Washington State University, Everett -- 05/2026",
        "A.A.S. in Computer Science, Everett Community College -- 04/2024 -- GPA: 4.0 (President's List)",
        "B.S. in Physics, University of Maryland, College Park -- 12/2015",
    ]

    certifications = [
        "Network Security+, CompTIA (07/2024)",
        "Cybersecurity Analyst, Everett Community College (07/2024)",
        "Data Analytics, Washington State University (12/2024)",
        "Medical Person in Charge, U.S. Coast Guard Certified",
    ]

    publication = "Journal of Instrumentation (2016) -- Liquid Scintillator Tiles for Calorimetry, Large Hadron Collider detector performance analysis"

    return ResumePayload(
        name=name,
        location=location,
        phone=phone,
        email=email,
        linkedin=linkedin,
        github=github,
        headline=headline,
        summary=summary,
        skills=skills,
        experience=experience,
        projects=projects,
        education=education,
        certifications=certifications,
        publication=publication,
    )


def latex_list(items: Iterable[str]) -> str:
    return "\\begin{itemize}[leftmargin=0.18in, itemsep=0.2em]\n" + "\n".join(
        f"  \\item {latex_escape(normalize_dash(item))}" for item in items if item.strip()
    ) + "\n\\end{itemize}"


def split_name(full_name: str) -> Tuple[str, str]:
    cleaned = full_name.strip()
    if not cleaned:
        return "Jeff", "Calderon"
    parts = cleaned.split()
    if len(parts) == 1:
        return parts[0], ""
    return parts[0], " ".join(parts[1:])


def split_location(location: str) -> Tuple[str, str]:
    cleaned = location.strip()
    if not cleaned:
        return "Marysville", "WA"
    if "," in cleaned:
        city, state = cleaned.split(",", 1)
        return city.strip(), state.strip()
    return cleaned, ""


def replace_pattern(text: str, pattern: str, replacement: str) -> str:
    return re.sub(pattern, lambda _match: replacement, text, count=1, flags=re.MULTILINE)


def replace_between_markers(text: str, start_marker: str, end_marker: str, replacement: str) -> str:
    start = text.find(start_marker)
    if start == -1:
        raise ValueError(f"Start marker not found: {start_marker}")
    end = text.find(end_marker, start)
    if end == -1:
        raise ValueError(f"End marker not found: {end_marker}")
    return text[:start] + replacement + text[end:]


def build_summary_block(resume: ResumePayload) -> str:
    return (
        "% TAGLINE / SUMMARY-------------------------------\n"
        "\\vspace*{-9mm}\n"
        "\\begin{center}\n"
        f"\\textit{{{latex_escape(resume.headline)}}}\n"
        "\\end{center}\n"
        "\\vspace*{1mm}\n\n"
        f"{latex_escape(resume.summary)}\n"
    )


def build_skill_block(resume: ResumePayload) -> str:
    lines: List[str] = []
    for label, values in resume.skills:
        if not values:
            continue
        lines.append(f"\\textbf{{{latex_escape(label)}:}} {latex_escape(', '.join(values))}")
    if not lines:
        lines = ["\\textbf{Skills:} None provided"]
    return (
        "% TECHNICAL SKILLS-------------------------------\n"
        "\\vspace*{1mm}\n"
        "\\section{Technical Skills}\n"
        "\\vspace*{1mm}\n"
        + "\n\n".join(lines)
        + "\n"
    )


def render_template(template_text: str, resume: ResumePayload) -> str:
    rendered = template_text
    first_name, last_name = split_name(resume.name)
    city, state = split_location(resume.location)
    rendered = replace_pattern(rendered, r"^\\name\{.*?\}\{.*?\}$", f"\\name{{{latex_escape(first_name)}}}{{{latex_escape(last_name)}}}")
    rendered = replace_pattern(rendered, r"^\\address\{.*?\}\{.*?\}$", f"\\address{{{latex_escape(city)}}}{{{latex_escape(state)}}}")
    rendered = replace_pattern(rendered, r"^\\phone\{.*?\}$", f"\\phone{{{latex_escape(resume.phone)}}}")
    rendered = replace_pattern(rendered, r"^\\email\{.*?\}$", f"\\email{{{latex_escape(resume.email)}}}")
    rendered = replace_pattern(rendered, r"^\\social\[linkedin\]\{.*?\}$", f"\\social[linkedin]{{{latex_escape(resume.linkedin.replace('https://', '').replace('http://', '').replace('www.', ''))}}}")
    rendered = replace_pattern(rendered, r"^\\social\[github\]\{.*?\}$", f"\\social[github]{{{latex_escape(resume.github.replace('https://', '').replace('http://', '').replace('www.', ''))}}}")

    rendered = replace_between_markers(rendered, "% TAGLINE / SUMMARY-------------------------------", "% TECHNICAL SKILLS-------------------------------", build_summary_block(resume))
    rendered = replace_between_markers(rendered, "% TECHNICAL SKILLS-------------------------------", "% PROFESSIONAL EXPERIENCE-------------------------------", build_skill_block(resume))
    rendered = replace_between_markers(rendered, "% PROFESSIONAL EXPERIENCE-------------------------------", "% ML AND ANALYTICS PROJECTS-------------------------------", generate_experience_block(resume))
    rendered = replace_between_markers(rendered, "% ML AND ANALYTICS PROJECTS-------------------------------", "% EDUCATION-------------------------------", generate_projects_block(resume))
    rendered = replace_between_markers(rendered, "% EDUCATION-------------------------------", "% CERTIFICATIONS-------------------------------", generate_education_block(resume))
    rendered = replace_between_markers(rendered, "% CERTIFICATIONS-------------------------------", "% PUBLICATION-------------------------------", generate_certification_block(resume))
    rendered = replace_between_markers(rendered, "% PUBLICATION-------------------------------", "\\end{document}", generate_publication_block(resume))
    return rendered


def bullet_block(lines: Sequence[str]) -> str:
    items = "\n".join(f"  \\item {latex_escape(normalize_dash(line))}" for line in lines)
    return f"\\begin{{itemize}}[label={{-}}, itemsep=-1.0ex]\n{items}\n\\end{{itemize}}"


def generate_section_body(section: ResumeSection) -> str:
    """Render a section entry with optional right-aligned date."""
    title = latex_escape(section.title)
    if section.date_string:
        return f"\\textbf{{{title}}} \\hfill {latex_escape(section.date_string)}\n\\vspace*{{-2mm}}\n{bullet_block(section.body_lines)}\n"
    return f"\\textbf{{{title}}}\n\\vspace*{{-2mm}}\n{bullet_block(section.body_lines)}\n"


def generate_experience_block(resume: ResumePayload) -> str:
    blocks = []
    blocks.append("% PROFESSIONAL EXPERIENCE-------------------------------\n\\vspace*{2mm}\n\\section{Professional Experience}\n\\vspace*{1mm}\n")
    for section in resume.experience:
        blocks.append(generate_section_body(section))
    return "".join(blocks)


def generate_projects_block(resume: ResumePayload) -> str:
    if not resume.projects:
        return ""
    blocks = []
    blocks.append("% ML AND ANALYTICS PROJECTS-------------------------------\n\\vspace*{2mm}\n\\section{ML and Analytics Projects}\n\\vspace*{1mm}\n")
    for section in resume.projects:
        blocks.append(generate_section_body(section))
    return "".join(blocks)


def generate_education_block(resume: ResumePayload) -> str:
    if not resume.education:
        return ""
    lines = [latex_escape(line) for line in resume.education]
    body = "\\\\\n".join(lines)
    return f"% EDUCATION-------------------------------\n\\vspace*{{2mm}}\n\\section{{Education}}\n\\vspace*{{1mm}}\n{body}\n"


def generate_certification_block(resume: ResumePayload) -> str:
    if not resume.certifications:
        return ""
    lines = [latex_escape(line) for line in resume.certifications]
    body = "\\\\\n".join(lines)
    return f"% CERTIFICATIONS-------------------------------\n\\vspace*{{2mm}}\n\\section{{Certifications}}\n\\vspace*{{1mm}}\n{body}\n"


def generate_publication_block(resume: ResumePayload) -> str:
    if not resume.publication:
        return ""
    return (
        "% PUBLICATION-------------------------------\n"
        "\\vspace*{2mm}\n\\section{Publication}\n\\vspace*{1mm}\n"
        f"\\textit{{Journal of Instrumentation}} (2016) -- Liquid Scintillator Tiles for Calorimetry, Large Hadron Collider detector performance analysis\n"
    )


def compile_latex(tex_path: Path, working_dir: Path) -> Path:
    """Compile a .tex file to PDF using pdflatex (from TeX Live).

    Runs pdflatex twice for proper cross-references and ToC resolution.
    Uses -interaction=nonstopmode to avoid hanging on errors.
    """
    command = [
        "pdflatex",
        "-interaction=nonstopmode",
        "-output-directory", str(working_dir),
        str(tex_path),
    ]
    # First pass: generate .aux files
    subprocess.run(command, cwd=working_dir, check=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    # Second pass: resolve references
    subprocess.run(command, cwd=working_dir, check=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    pdf_path = tex_path.with_suffix(".pdf")
    if not pdf_path.exists():
        # Check .log for the error
        log_path = tex_path.with_suffix(".log")
        if log_path.exists():
            log_tail = log_path.read_text(errors="ignore")[-2000:]
            raise FileNotFoundError(f"pdflatex failed to produce PDF. Log tail:\n{log_tail}")
        raise FileNotFoundError(f"Expected PDF not produced: {pdf_path}")
    return pdf_path


def build_docx(resume: ResumePayload, output_dir: Path, safe_title: str) -> Path:
    """Generate a DOCX that visually matches the moderncv banking PDF style."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor as DocxRGB, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn as docx_qn
    
    doc = DocxDocument()
    
    # Page setup matching the PDF (Letter: 8.5x11 inches, with 0.3in margins)
    sec = doc.sections[0]
    # Letter dimensions in EMU: 8.5 x 914400 = 7772400, 11 x 914400 = 10058400
    sec.page_width = Emu(7772400)
    sec.page_height = Emu(10058400)
    # 0.3 inch margins = 0.3 x 914400 = 274320 EMU
    sec.left_margin = Emu(274320)
    sec.right_margin = Emu(274320)
    sec.top_margin = Emu(274320)
    sec.bottom_margin = Emu(274320)
    
    FONT = "Arial"
    COLOR_ACCENT = DocxRGB(0x1F, 0x4E, 0x79)  # dark blue, matches template
    COLOR_BODY = DocxRGB(0x1A, 0x1A, 0x1A)
    COLOR_MUTED = DocxRGB(0x44, 0x44, 0x44)
    
    def set_run(run, size=10, bold=False, italic=False, color=COLOR_BODY):
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(docx_qn('w:rFonts'))
        if rFonts is None:
            rFonts = run._element.makeelement(docx_qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(docx_qn('w:ascii'), FONT)
        rFonts.set(docx_qn('w:hAnsi'), FONT)
    
    def add_para(sp_before=0, sp_after=2, alignment=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(sp_before)
        p.paragraph_format.space_after = Pt(sp_after)
        if alignment is not None:
            p.alignment = alignment
        r = p.add_run("")
        return p, r
    
    # ─── HEADER (centered, like moderncv banking) ───
    p, r = add_para(sp_before=0, sp_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r.text = resume.name
    set_run(r, size=16, bold=True, color=COLOR_ACCENT)
    
    # Location line (centered, under name)
    p, r = add_para(sp_before=0, sp_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r.text = resume.location
    set_run(r, size=9, color=COLOR_MUTED)
    
    # Contact line (centered)
    p, r = add_para(sp_before=2, sp_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r.text = f"{resume.phone}  |  {resume.email}  |  {resume.linkedin}  |  {resume.github}"
    set_run(r, size=8, color=COLOR_MUTED)
    
    # ─── SUMMARY ───
    if resume.summary:
        p, r = add_para(sp_before=4, sp_after=4)
        r.text = resume.summary
        set_run(r, size=8)
    
    # ─── TECHNICAL SKILLS ───
    if resume.skills:
        p, r = add_para(sp_before=8, sp_after=1)
        r.text = "Technical Skills"
        set_run(r, size=12, bold=True, color=COLOR_ACCENT)
        # Add bottom border to match moderncv section rule
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(docx_qn('w:pBdr'), {})
        bottom = pBdr.makeelement(docx_qn('w:bottom'), {
            docx_qn('w:val'): 'single',
            docx_qn('w:sz'): '6',
            docx_qn('w:space'): '1',
            docx_qn('w:color'): '1F4E79'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for label, values in resume.skills:
            if not values:
                continue
            p, r = add_para(sp_before=1, sp_after=1)
            r.text = f"{label}: {', '.join(values)}"
            set_run(r, size=9)
    
    # ─── PROFESSIONAL EXPERIENCE ───
    if resume.experience:
        p, r = add_para(sp_before=8, sp_after=1)
        r.text = "Professional Experience"
        set_run(r, size=12, bold=True, color=COLOR_ACCENT)
        # Add bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(docx_qn('w:pBdr'), {})
        bottom = pBdr.makeelement(docx_qn('w:bottom'), {
            docx_qn('w:val'): 'single',
            docx_qn('w:sz'): '6',
            docx_qn('w:space'): '1',
            docx_qn('w:color'): '1F4E79'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for section in resume.experience:
            p, r = add_para(sp_before=4, sp_after=0)
            r.text = section.title
            set_run(r, size=9, bold=True)
            
            if section.date_string:
                run2 = p.add_run(f"    {section.date_string}")
                set_run(run2, size=9, italic=True, color=COLOR_MUTED)
            
            for bullet in section.body_lines:
                p, r = add_para(sp_before=0, sp_after=0)
                r.text = f"- {bullet}"
                set_run(r, size=9)
    
    # ─── PROJECTS ───
    if resume.projects:
        p, r = add_para(sp_before=8, sp_after=1)
        r.text = "Projects"
        set_run(r, size=12, bold=True, color=COLOR_ACCENT)
        # Add bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(docx_qn('w:pBdr'), {})
        bottom = pBdr.makeelement(docx_qn('w:bottom'), {
            docx_qn('w:val'): 'single',
            docx_qn('w:sz'): '6',
            docx_qn('w:space'): '1',
            docx_qn('w:color'): '1F4E79'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for section in resume.projects:
            p, r = add_para(sp_before=3, sp_after=0)
            r.text = section.title
            set_run(r, size=9, bold=True)
            
            if section.date_string:
                run2 = p.add_run(f"    {section.date_string}")
                set_run(run2, size=9, italic=True, color=COLOR_MUTED)
            
            for bullet in section.body_lines:
                p, r = add_para(sp_before=0, sp_after=0)
                r.text = f"- {bullet}"
                set_run(r, size=9)
    
    # ─── EDUCATION ───
    if resume.education:
        p, r = add_para(sp_before=8, sp_after=1)
        r.text = "Education"
        set_run(r, size=12, bold=True, color=COLOR_ACCENT)
        # Add bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(docx_qn('w:pBdr'), {})
        bottom = pBdr.makeelement(docx_qn('w:bottom'), {
            docx_qn('w:val'): 'single',
            docx_qn('w:sz'): '6',
            docx_qn('w:space'): '1',
            docx_qn('w:color'): '1F4E79'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for line in resume.education:
            p, r = add_para(sp_before=1, sp_after=1)
            r.text = line
            set_run(r, size=9)
    
    # ─── CERTIFICATIONS ───
    if resume.certifications:
        p, r = add_para(sp_before=8, sp_after=1)
        r.text = "Certifications"
        set_run(r, size=12, bold=True, color=COLOR_ACCENT)
        # Add bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(docx_qn('w:pBdr'), {})
        bottom = pBdr.makeelement(docx_qn('w:bottom'), {
            docx_qn('w:val'): 'single',
            docx_qn('w:sz'): '6',
            docx_qn('w:space'): '1',
            docx_qn('w:color'): '1F4E79'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for line in resume.certifications:
            p, r = add_para(sp_before=1, sp_after=1)
            r.text = line
            set_run(r, size=9)
    
    docx_path = output_dir / f"{safe_title}.docx"
    doc.save(str(docx_path))
    return docx_path


def build_resume_from_payload(payload: Dict[str, Any], output_dir: Optional[Path] = None, job_title: Optional[str] = None) -> Tuple[Path, Path]:
    """Build PDF + DOCX from a structured JSON payload (the Imaginator output schema).
    
    Args:
        payload: Dict with name, location, phone, email, linkedin, github,
                 headline, summary, skills, experience, projects, education, certifications
        output_dir: Where to write generated files (defaults to temp dir)
        job_title: Optional headline override
    
    Returns:
        (pdf_path, docx_path) tuple
    """
    template_path = Path(__file__).parent / "main_template.tex"
    if output_dir is None:
        output_dir = DEFAULT_OUTPUT_DIR
    output_dir.mkdir(parents=True, exist_ok=True)

    resume = parse_resume_payload({"data": payload} if "data" not in payload else payload)
    if job_title:
        resume.headline = job_title

    template_text = template_path.read_text(encoding="utf-8")
    rendered_tex = render_template(template_text, resume)

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", job_title or resume.headline or "resume")
    tex_path = output_dir / f"{safe_title}.tex"
    tex_path.write_text(rendered_tex, encoding="utf-8")

    pdf_path = compile_latex(tex_path, output_dir)
    docx_path = build_docx(resume, output_dir, safe_title)
    return pdf_path, docx_path


def build_resume_from_payload_base64(payload: Dict[str, Any], job_title: Optional[str] = None) -> Dict[str, str]:
    """Build PDF + DOCX and return as base64-encoded strings.
    
    Returns:
        Dict with 'pdf_base64', 'docx_base64', 'filename' keys
    """
    import base64
    output_dir = DEFAULT_OUTPUT_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    
    pdf_path, docx_path = build_resume_from_payload(payload, output_dir, job_title)
    
    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", job_title or "resume")
    
    with open(pdf_path, "rb") as f:
        pdf_b64 = base64.b64encode(f.read()).decode("ascii")
    with open(docx_path, "rb") as f:
        docx_b64 = base64.b64encode(f.read()).decode("ascii")
    
    return {
        "pdf_base64": pdf_b64,
        "docx_base64": docx_b64,
        "filename": safe_title,
    }


def build_resume(input_path: Path, template_path: Path, output_dir: Path, job_title: Optional[str] = None) -> Tuple[Path, Path]:
    """Build both PDF and DOCX from the input. Returns (pdf_path, docx_path)."""
    output_dir.mkdir(parents=True, exist_ok=True)

    if input_path.suffix.lower() in {".json"}:
        bundle = json.loads(input_path.read_text(encoding="utf-8"))
    else:
        bundle = extract_text_bundle(input_path)

    resume = parse_resume_payload(bundle)
    if job_title:
        resume.headline = job_title

    template_text = template_path.read_text(encoding="utf-8")
    rendered_tex = render_template(template_text, resume)

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", job_title or resume.headline or "resume")
    tex_path = output_dir / f"{safe_title}.tex"
    tex_path.write_text(rendered_tex, encoding="utf-8")

    pdf_path = compile_latex(tex_path, output_dir)
    docx_path = build_docx(resume, output_dir, safe_title)
    return pdf_path, docx_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a resume PDF from an Imaginator export bundle.")
    parser.add_argument("input", type=Path, help="Path to the export_results bundle text file or JSON file")
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE, help="LaTeX template to populate")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR, help="Directory for generated .tex and .pdf files")
    parser.add_argument("--job-title", type=str, default=None, help="Optional headline/job title override")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    pdf_path, docx_path = build_resume(args.input, args.template, args.output_dir, args.job_title)
    print(f"Generated PDF: {pdf_path}")
    print(f"Generated DOCX: {docx_path}")


if __name__ == "__main__":
    main()